import http.server
import socketserver
import os
import urllib.request
import urllib.error
import json
import mimetypes
import shutil
import urllib.parse
import re

PORT = 8000

def is_system_or_dangerous_path(path):
    if not path:
        return True
    
    try:
        path = os.path.abspath(path)
    except Exception:
        return True
        
    if path in ('/', '\\') or re.match(r'^[a-zA-Z]:\\?$', path):
        return True
        
    dangerous_dirs = [
        '/boot', '/dev', '/etc', '/lib', '/lib64', '/proc', '/sys', '/bin', '/sbin', '/usr', '/var', '/root',
        'c:\\windows', 'c:\\program files', 'c:\\program files (x86)', 'c:\\users\\public', 'c:\\recovery',
        'c:\\system volume information', 'c:\\windows.old'
    ]
    
    normalized = path.lower().replace('/', '\\') if os.name == 'nt' else path.lower()
    
    for d in dangerous_dirs:
        norm_d = d.lower().replace('/', '\\') if os.name == 'nt' else d.lower()
        if normalized == norm_d or normalized.startswith(norm_d + os.sep):
            return True
            
    parts = list(filter(None, path.split(os.sep)))
    
    if os.name != 'nt':
        if len(parts) <= 2:
            return True
    else:
        if len(parts) >= 1 and parts[0].endswith(':'):
            if len(parts) <= 3:
                return True
        else:
            if len(parts) <= 2:
                return True
                
    return False

def set_file_permissions_and_ownership(file_path):
    try:
        os.chmod(file_path, 0o666)
    except Exception:
        pass
    try:
        parent_dir = os.path.dirname(file_path)
        if parent_dir and os.path.exists(parent_dir):
            stat_info = os.stat(parent_dir)
            os.chown(file_path, stat_info.st_uid, stat_info.st_gid)
    except Exception:
        pass

class NoCacheProxyHTTPRequestHandler(http.server.SimpleHTTPRequestHandler):
    protocol_version = "HTTP/1.1"

    def end_headers(self):
        self.send_header('Cache-Control', 'no-store, no-cache, must-revalidate, max-age=0')
        self.send_header('Pragma', 'no-cache')
        self.send_header('Expires', '0')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, X-Ollama-Host, Accept')
        self.send_header('Connection', 'close')
        self.close_connection = True
        super().end_headers()

    def do_OPTIONS(self):
        self.send_response(204)
        self.end_headers()

    def do_GET(self):
        if self.path.startswith('/ollama-api/'):
            # Proxy GET request to Ollama
            target_path = self.path[12:] # strip '/ollama-api/'
            ollama_host = self.headers.get('X-Ollama-Host', 'http://localhost:11434').rstrip('/')
            target_url = f"{ollama_host}/{target_path}"
            
            try:
                req = urllib.request.Request(target_url)
                # Copy headers from request
                for header in ['Accept', 'Content-Type']:
                    if header in self.headers:
                        req.add_header(header, self.headers[header])
                
                with urllib.request.urlopen(req) as response:
                    self.send_response(response.status)
                    # Copy headers from response
                    for header, val in response.getheaders():
                        if header.lower() not in [
                            'access-control-allow-origin', 
                            'access-control-allow-methods', 
                            'access-control-allow-headers', 
                            'access-control-expose-headers',
                            'content-length', 
                            'transfer-encoding',
                            'connection'
                        ]:
                            self.send_header(header, val)
                    self.end_headers()
                    self.wfile.write(response.read())
            except urllib.error.HTTPError as e:
                self.send_response(e.code)
                for header, val in e.headers.items():
                    if header.lower() not in [
                        'access-control-allow-origin', 
                        'access-control-allow-methods', 
                        'access-control-allow-headers', 
                        'access-control-expose-headers',
                        'content-length', 
                        'transfer-encoding',
                        'connection'
                    ]:
                        self.send_header(header, val)
                self.end_headers()
                self.wfile.write(e.read())
            except urllib.error.URLError as e:
                self.send_response(502)
                self.send_header('Content-Type', 'text/plain; charset=utf-8')
                self.end_headers()
                self.wfile.write(f"Ollama server connection failed: {e.reason}".encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'text/plain; charset=utf-8')
                self.end_headers()
                self.wfile.write(f"Proxy error: {str(e)}".encode('utf-8'))
        elif self.path.startswith('/api/folder'):
            # Parse GET query parameter 'path'
            parsed_url = urllib.parse.urlparse(self.path)
            query_params = urllib.parse.parse_qs(parsed_url.query)
            folder_path = query_params.get('path', [None])[0]
            
            if not folder_path or not os.path.exists(folder_path):
                self.send_response(400)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": "Путь не найден или не указан"}).encode('utf-8'))
                return
                
            try:
                files_list = []
                for entry in os.scandir(folder_path):
                    if entry.name.startswith('.'):
                        continue
                    
                    stat = entry.stat()
                    is_dir = entry.is_dir()
                    mime_type, _ = mimetypes.guess_type(entry.path)
                    ext = os.path.splitext(entry.name)[1].lower()
                    
                    files_list.append({
                        "name": entry.name,
                        "isDir": is_dir,
                        "size": stat.st_size if not is_dir else 0,
                        "ext": ext,
                        "mimeType": mime_type or ("directory" if is_dir else "application/octet-stream"),
                        "modified": int(stat.st_mtime * 1000)
                    })
                
                files_list.sort(key=lambda x: (not x['isDir'], x['name'].lower()))
                
                self.send_response(200)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"path": folder_path, "files": files_list}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": f"Ошибка чтения папки: {str(e)}"}).encode('utf-8'))
        elif self.path.startswith('/api/file/read'):
            # Parse GET query parameter 'path'
            parsed_url = urllib.parse.urlparse(self.path)
            query_params = urllib.parse.parse_qs(parsed_url.query)
            file_path = query_params.get('path', [None])[0]
            
            if not file_path or not os.path.exists(file_path) or not os.path.isfile(file_path):
                self.send_response(400)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": "Файл не найден"}).encode('utf-8'))
                return
                
            if is_system_or_dangerous_path(os.path.dirname(file_path)):
                self.send_response(403)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": "Файл находится в системной или корневой папке. Доступ заблокирован."}).encode('utf-8'))
                return
                
            try:
                ext = os.path.splitext(file_path)[1].lower()
                content = ""
                if ext == '.docx':
                    try:
                        import docx
                        doc = docx.Document(file_path)
                        content = "\n".join([p.text for p in doc.paragraphs])
                    except Exception as docx_err:
                        self.send_response(500)
                        self.send_header('Content-Type', 'application/json; charset=utf-8')
                        self.end_headers()
                        self.wfile.write(json.dumps({"error": f"Не удалось прочитать .docx: {str(docx_err)}"}).encode('utf-8'))
                        return
                else:
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    except UnicodeDecodeError:
                        with open(file_path, 'r', encoding='cp1251') as f:
                            content = f.read()
                            
                self.send_response(200)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "content": content}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": f"Ошибка чтения файла: {str(e)}"}).encode('utf-8'))
        else:
            super().do_GET()

    def do_POST(self):
        if self.path.startswith('/ollama-api/'):
            # Proxy POST request to Ollama
            target_path = self.path[12:] # strip '/ollama-api/'
            ollama_host = self.headers.get('X-Ollama-Host', 'http://localhost:11434').rstrip('/')
            target_url = f"{ollama_host}/{target_path}"
            
            content_length = int(self.headers.get('Content-Length', 0))
            post_data = self.rfile.read(content_length) if content_length > 0 else None
            
            try:
                req = urllib.request.Request(target_url, data=post_data, method='POST')
                # Copy headers from request
                for header in ['Accept', 'Content-Type']:
                    if header in self.headers:
                        req.add_header(header, self.headers[header])
                
                with urllib.request.urlopen(req) as response:
                    self.send_response(response.status)
                    # Copy headers from response
                    for header, val in response.getheaders():
                        if header.lower() not in [
                            'access-control-allow-origin', 
                            'access-control-allow-methods', 
                            'access-control-allow-headers', 
                            'access-control-expose-headers',
                            'content-length', 
                            'transfer-encoding',
                            'connection'
                        ]:
                            self.send_header(header, val)
                    self.end_headers()
                    
                    # Stream response chunks back to client (essential for streaming effect)
                    while True:
                        chunk = response.readline()
                        if not chunk:
                            break
                        self.wfile.write(chunk)
                        self.wfile.flush()
            except urllib.error.HTTPError as e:
                self.send_response(e.code)
                for header, val in e.headers.items():
                    if header.lower() not in [
                        'access-control-allow-origin', 
                        'access-control-allow-methods', 
                        'access-control-allow-headers', 
                        'access-control-expose-headers',
                        'content-length', 
                        'transfer-encoding',
                        'connection'
                    ]:
                        self.send_header(header, val)
                self.end_headers()
                self.wfile.write(e.read())
            except urllib.error.URLError as e:
                self.send_response(502)
                self.send_header('Content-Type', 'text/plain; charset=utf-8')
                self.end_headers()
                self.wfile.write(f"Ollama server connection failed: {e.reason}".encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'text/plain; charset=utf-8')
                self.end_headers()
                self.wfile.write(f"Proxy error: {str(e)}".encode('utf-8'))
        elif self.path.startswith('/api/folder/locate'):
            # locate directory by name and sample files
            try:
                content_length = int(self.headers.get('Content-Length', 0))
                post_data = self.rfile.read(content_length).decode('utf-8')
                params = json.loads(post_data)
                folder_name = params.get('folderName')
                sample_files = params.get('sampleFiles', [])
                
                if not folder_name:
                    self.send_response(400)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Имя папки не указано"}).encode('utf-8'))
                    return
                
                search_root = "/home"
                if os.name == 'nt':
                    search_root = os.path.expanduser("~")
                    
                matches = []
                sample_set = set(sample_files)
                
                # Check search_root itself first
                if os.path.basename(search_root.rstrip(os.sep)) == folder_name:
                    try:
                        score = sum(1 for f in sample_set if os.path.exists(os.path.join(search_root, f)))
                        matches.append((search_root, score))
                    except Exception:
                        pass
                
                # Walk search_root
                for root, dirs, files_in_dir in os.walk(search_root):
                    # Prune large or hidden directories to ensure extremely fast search
                    dirs[:] = [d for d in dirs if not d.startswith('.') and d.lower() not in ('node_modules', 'venv', 'env', 'build', 'dist', 'target', 'cache', 'library', 'appdata')]
                    
                    for d in dirs:
                        if d == folder_name:
                            full_path = os.path.join(root, d)
                            try:
                                score = 0
                                actual_files = os.listdir(full_path)
                                for f in sample_set:
                                    if f in actual_files:
                                        score += 1
                                matches.append((full_path, score))
                            except Exception:
                                pass
                                
                if not matches:
                    self.send_response(404)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": f"Папка с именем '{folder_name}' не найдена на сервере"}).encode('utf-8'))
                    return
                    
                # Sort matches by score descending, then by path length
                matches.sort(key=lambda x: (-x[1], len(x[0])))
                best_match = matches[0][0]
                
                self.send_response(200)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "path": best_match}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": f"Ошибка поиска папки: {str(e)}"}).encode('utf-8'))
        elif self.path.startswith('/api/folder/sort'):
            # organize files into categories
            try:
                CATEGORY_MAP = {
                    # Documents
                    '.txt': 'documents', '.md': 'documents', '.pdf': 'documents', 
                    '.doc': 'documents', '.docx': 'documents', '.rtf': 'documents', '.odt': 'documents',
                    # Images
                    '.png': 'images', '.jpg': 'images', '.jpeg': 'images', '.gif': 'images', 
                    '.bmp': 'images', '.svg': 'images', '.webp': 'images',
                    # Code
                    '.py': 'code', '.js': 'code', '.html': 'code', '.css': 'code', 
                    '.json': 'code', '.c': 'code', '.cpp': 'code', '.h': 'code', 
                    '.java': 'code', '.go': 'code', '.sh': 'code', '.rs': 'code',
                    # Data
                    '.csv': 'data', '.xlsx': 'data', '.xls': 'data', '.xml': 'data', 
                    '.db': 'data', '.sql': 'data',
                    # Archives
                    '.zip': 'archives', '.tar': 'archives', '.gz': 'archives', 
                    '.rar': 'archives', '.7z': 'archives'
                }
                
                content_length = int(self.headers.get('Content-Length', 0))
                post_data = self.rfile.read(content_length).decode('utf-8')
                params = json.loads(post_data)
                folder_path = params.get('path')
                
                if not folder_path or not os.path.exists(folder_path):
                    self.send_response(400)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Путь не найден или не указан"}).encode('utf-8'))
                    return
                    
                if is_system_or_dangerous_path(folder_path):
                    self.send_response(403)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Выбранная папка является системной или корневой. Операции в ней запрещены из соображений безопасности."}).encode('utf-8'))
                    return
                    
                moved = []
                for entry in os.scandir(folder_path):
                    if entry.is_file() and not entry.name.startswith('.'):
                        ext = os.path.splitext(entry.name)[1].lower()
                        category = CATEGORY_MAP.get(ext, 'other')
                        
                        category_dir = os.path.join(folder_path, category)
                        os.makedirs(category_dir, exist_ok=True)
                        
                        src = entry.path
                        dst = os.path.join(category_dir, entry.name)
                        
                        final_dst = dst
                        if os.path.exists(dst):
                            base, ext_part = os.path.splitext(entry.name)
                            counter = 1
                            while os.path.exists(final_dst):
                                final_dst = os.path.join(category_dir, f"{base}_{counter}{ext_part}")
                                counter += 1
                                
                        shutil.move(src, final_dst)
                        moved.append({"name": entry.name, "from": src, "to": final_dst, "category": category})
                        
                self.send_response(200)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "moved": moved}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": f"Ошибка сортировки: {str(e)}"}).encode('utf-8'))
                
        elif self.path.startswith('/api/folder/rename'):
            # batch rename files in folder
            try:
                content_length = int(self.headers.get('Content-Length', 0))
                post_data = self.rfile.read(content_length).decode('utf-8')
                params = json.loads(post_data)
                folder_path = params.get('path')
                pattern = params.get('pattern', '')
                replace = params.get('replace', '')
                is_regex = params.get('regex', False)
                case_mode = params.get('caseMode', 'none') # 'none', 'upper', 'lower'
                prefix = params.get('prefix', '')
                suffix = params.get('suffix', '')
                
                if not folder_path or not os.path.exists(folder_path):
                    self.send_response(400)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Путь не найден или не указан"}).encode('utf-8'))
                    return
                    
                if is_system_or_dangerous_path(folder_path):
                    self.send_response(403)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Выбранная папка является системной или корневой. Операции в ней запрещены из соображений безопасности."}).encode('utf-8'))
                    return
                    
                renamed = []
                for entry in os.scandir(folder_path):
                    if entry.is_file() and not entry.name.startswith('.'):
                        old_name = entry.name
                        base, ext = os.path.splitext(old_name)
                        new_base = base
                        
                        if pattern:
                            if is_regex:
                                try:
                                    if re.search(pattern, old_name):
                                        temp_name = re.sub(pattern, replace, old_name)
                                        new_base, new_ext = os.path.splitext(temp_name)
                                        ext = new_ext
                                    else:
                                        new_base = base
                                except Exception as re_err:
                                    self.send_response(400)
                                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                                    self.end_headers()
                                    self.wfile.write(json.dumps({"error": f"Некорректное регулярное выражение: {str(re_err)}"}).encode('utf-8'))
                                    return
                            else:
                                if pattern in old_name:
                                    temp_name = old_name.replace(pattern, replace)
                                    new_base, new_ext = os.path.splitext(temp_name)
                                    ext = new_ext
                                else:
                                    new_base = base
                                
                        if case_mode == 'upper':
                            new_base = new_base.upper()
                        elif case_mode == 'lower':
                            new_base = new_base.lower()
                            
                        if prefix:
                            new_base = prefix + new_base
                        if suffix:
                            new_base = new_base + suffix
                            
                        new_name = new_base + ext
                        
                        if old_name != new_name:
                            src = entry.path
                            dst = os.path.join(folder_path, new_name)
                            
                            if os.path.exists(dst):
                                counter = 1
                                while os.path.exists(dst):
                                    dst = os.path.join(folder_path, f"{new_base}_{counter}{ext}")
                                    counter += 1
                                    
                            os.rename(src, dst)
                            renamed.append({"oldName": old_name, "newName": os.path.basename(dst)})
                            
                self.send_response(200)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "renamed": renamed}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": f"Ошибка переименования: {str(e)}"}).encode('utf-8'))
                
        elif self.path.startswith('/api/folder/delete'):
            # batch delete files in folder
            try:
                content_length = int(self.headers.get('Content-Length', 0))
                post_data = self.rfile.read(content_length).decode('utf-8')
                params = json.loads(post_data)
                folder_path = params.get('path')
                files_to_delete = params.get('files', [])
                
                if not folder_path or not os.path.exists(folder_path):
                    self.send_response(400)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Путь не найден или не указан"}).encode('utf-8'))
                    return
                    
                if is_system_or_dangerous_path(folder_path):
                    self.send_response(403)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Выбранная папка является системной или корневой. Операции в ней запрещены из соображений безопасности."}).encode('utf-8'))
                    return
                    
                deleted = []
                for file_name in files_to_delete:
                    safe_name = os.path.basename(file_name)
                    file_path = os.path.join(folder_path, safe_name)
                    if os.path.exists(file_path) and os.path.isfile(file_path):
                        os.remove(file_path)
                        deleted.append(safe_name)
                        
                self.send_response(200)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "deleted": deleted}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": f"Ошибка удаления: {str(e)}"}).encode('utf-8'))
                
        elif self.path.startswith('/api/folder/format_docx'):
            # format all docx files in folder according to GOST
            try:
                import docx
                from docx.shared import Pt, Cm
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": "Библиотека python-docx не установлена на сервере. Пожалуйста, убедитесь, что она установлена."}).encode('utf-8'))
                return

            try:
                content_length = int(self.headers.get('Content-Length', 0))
                post_data = self.rfile.read(content_length).decode('utf-8')
                params = json.loads(post_data)
                folder_path = params.get('path')
                
                if not folder_path or not os.path.exists(folder_path):
                    self.send_response(400)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Путь не найден или не указан"}).encode('utf-8'))
                    return
                    
                if is_system_or_dangerous_path(folder_path):
                    self.send_response(403)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Выбранная папка является системной или корневой. Операции в ней запрещены из соображений безопасности."}).encode('utf-8'))
                    return
                    
                formatted_files = []
                for entry in os.scandir(folder_path):
                    if entry.is_file() and entry.name.lower().endswith('.docx') and not entry.name.startswith('~$'):
                        file_path = entry.path
                        
                        # Apply GOST to docx
                        doc = docx.Document(file_path)
                        
                        # Margins
                        for section in doc.sections:
                            section.top_margin = Cm(2.0)
                            section.bottom_margin = Cm(2.0)
                            section.left_margin = Cm(3.0)
                            section.right_margin = Cm(1.5)
                            
                        # Paragraphs
                        for p in doc.paragraphs:
                            # 1. Clean quotes & dashes
                            text = p.text
                            text = re.sub(r'(^|[\s\(\[\{])"', r'\1«', text)
                            text = re.sub(r'"($|[\s\)\}\]\.,;:!\?])', r'»\1', text)
                            parts = text.split('"')
                            alt_parts = []
                            for i, part in enumerate(parts):
                                alt_parts.append(part)
                                if i < len(parts) - 1:
                                    alt_parts.append('«' if i % 2 == 0 else '»')
                            text = "".join(alt_parts)
                            text = re.sub(r'\s+-\s+', ' — ', text)
                            text = re.sub(r' {2,}', ' ', text)
                            p.text = text
                            
                            # Alignment and spacing
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p.paragraph_format.line_spacing = 1.5
                            if text.strip():
                                p.paragraph_format.first_line_indent = Cm(1.25)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.space_before = Pt(0)
                            
                            # Font
                            for run in p.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(14)
                                
                        doc.save(file_path)
                        set_file_permissions_and_ownership(file_path)
                        formatted_files.append(entry.name)
                        
                self.send_response(200)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "formatted": formatted_files}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": f"Ошибка форматирования ворд-файлов: {str(e)}"}).encode('utf-8'))

        elif self.path.startswith('/api/file/create'):
            # create and populate file (txt, md, or docx)
            try:
                content_length = int(self.headers.get('Content-Length', 0))
                post_data = self.rfile.read(content_length).decode('utf-8')
                params = json.loads(post_data)
                folder_path = params.get('path')
                filename = params.get('filename')
                content = params.get('content', '')
                
                if not folder_path or not filename:
                    self.send_response(400)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Путь или имя файла не указаны"}).encode('utf-8'))
                    return
                    
                if is_system_or_dangerous_path(folder_path):
                    self.send_response(403)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Выбранная папка является системной или корневой. Операции в ней заблокированы из соображений безопасности."}).encode('utf-8'))
                    return
                    
                # Prevent directory traversal in filename
                safe_filename = os.path.basename(filename)
                file_path = os.path.join(folder_path, safe_filename)
                
                if safe_filename.lower().endswith('.docx'):
                    # Create Word document
                    try:
                        import docx
                        from docx.shared import Pt, Cm
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                    except ImportError:
                        self.send_response(500)
                        self.send_header('Content-Type', 'application/json; charset=utf-8')
                        self.end_headers()
                        self.wfile.write(json.dumps({"error": "Библиотека python-docx не установлена на сервере. Не удалось создать файл Word (.docx)"}).encode('utf-8'))
                        return
                        
                    doc = docx.Document()
                    
                    # Margins
                    for section in doc.sections:
                        section.top_margin = Cm(2.0)
                        section.bottom_margin = Cm(2.0)
                        section.left_margin = Cm(3.0)
                        section.right_margin = Cm(1.5)
                        
                    # Add content paragraph by paragraph
                    lines = content.splitlines()
                    if not lines:
                        lines = ['']
                    for line in lines:
                        cleaned_line = line.strip()
                        p = doc.add_paragraph(cleaned_line)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.line_spacing = 1.5
                        if cleaned_line:
                            p.paragraph_format.first_line_indent = Cm(1.25)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(0)
                        
                        # Add run style
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(14)
                            
                    doc.save(file_path)
                    set_file_permissions_and_ownership(file_path)
                else:
                    # Create standard text file (txt, md, py, etc)
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    set_file_permissions_and_ownership(file_path)
                    
                # Delete old file if renaming
                old_filename = params.get('old_filename')
                if old_filename:
                    safe_old = os.path.basename(old_filename)
                    if safe_old != safe_filename:
                        old_path = os.path.join(folder_path, safe_old)
                        if os.path.exists(old_path) and os.path.isfile(old_path):
                            os.remove(old_path)
                        
                self.send_response(200)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "file": safe_filename}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": f"Ошибка создания файла: {str(e)}"}).encode('utf-8'))
                
        elif self.path.startswith('/api/file/format'):
            # format single text file (GOST or normalization)
            try:
                content_length = int(self.headers.get('Content-Length', 0))
                post_data = self.rfile.read(content_length).decode('utf-8')
                params = json.loads(post_data)
                file_path = params.get('filePath')
                rule = params.get('rule', 'gost')
                
                if not file_path or not os.path.exists(file_path) or not os.path.isfile(file_path):
                    self.send_response(400)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Файл не найден"}).encode('utf-8'))
                    return
                    
                if is_system_or_dangerous_path(os.path.dirname(file_path)):
                    self.send_response(403)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": "Файл находится в системной или корневой папке. Операции в ней запрещены из соображений безопасности."}).encode('utf-8'))
                    return
                    
                ext = os.path.splitext(file_path)[1].lower()
                if ext == '.docx':
                    if rule != 'gost':
                        self.send_response(400)
                        self.send_header('Content-Type', 'application/json; charset=utf-8')
                        self.end_headers()
                        self.wfile.write(json.dumps({"error": "Нормализация поддерживается только для текстовых файлов (.txt, .md)"}).encode('utf-8'))
                        return
                    
                    try:
                        import docx
                        from docx.shared import Pt, Cm
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                    except ImportError:
                        self.send_response(500)
                        self.send_header('Content-Type', 'application/json; charset=utf-8')
                        self.end_headers()
                        self.wfile.write(json.dumps({"error": "Библиотека python-docx не установлена на сервере. Не удалось отформатировать .docx"}).encode('utf-8'))
                        return
                    
                    try:
                        doc = docx.Document(file_path)
                        for section in doc.sections:
                            section.top_margin = Cm(2.0)
                            section.bottom_margin = Cm(2.0)
                            section.left_margin = Cm(3.0)
                            section.right_margin = Cm(1.5)
                            
                        for p in doc.paragraphs:
                            text = p.text
                            text = re.sub(r'(^|[\s\(\[\{])"', r'\1«', text)
                            text = re.sub(r'"($|[\s\)\}\]\.,;:!\?])', r'»\1', text)
                            parts = text.split('"')
                            alt_parts = []
                            for i, part in enumerate(parts):
                                alt_parts.append(part)
                                if i < len(parts) - 1:
                                    alt_parts.append('«' if i % 2 == 0 else '»')
                            text = "".join(alt_parts)
                            text = re.sub(r'\s+-\s+', ' — ', text)
                            text = re.sub(r' {2,}', ' ', text)
                            p.text = text
                            
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p.paragraph_format.line_spacing = 1.5
                            if text.strip():
                                p.paragraph_format.first_line_indent = Cm(1.25)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.space_before = Pt(0)
                            
                            for run in p.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(14)
                                
                        doc.save(file_path)
                        set_file_permissions_and_ownership(file_path)
                        self.send_response(200)
                        self.send_header('Content-Type', 'application/json; charset=utf-8')
                        self.end_headers()
                        self.wfile.write(json.dumps({"status": "success", "message": "Файл .docx успешно отформатирован по ГОСТ"}).encode('utf-8'))
                        return
                    except Exception as docx_err:
                        self.send_response(500)
                        self.send_header('Content-Type', 'application/json; charset=utf-8')
                        self.end_headers()
                        self.wfile.write(json.dumps({"error": f"Ошибка форматирования .docx: {str(docx_err)}"}).encode('utf-8'))
                        return
                
                content = ""
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    with open(file_path, 'r', encoding='cp1251') as f:
                        content = f.read()
                        
                formatted_content = content
                
                if rule == 'normalize':
                    lines = content.splitlines()
                    cleaned_lines = []
                    for line in lines:
                        line_converted = line.replace('\t', '    ')
                        cleaned_lines.append(line_converted.rstrip(' '))
                    formatted_content = '\n'.join(cleaned_lines)
                    
                elif rule == 'gost':
                    # 1. straight quotes to Russian guillemets
                    formatted_content = re.sub(r'(^|[\s\(\[\{])"', r'\1«', formatted_content)
                    formatted_content = re.sub(r'"($|[\s\)\}\]\.,;:!\?])', r'»\1', formatted_content)
                    parts = formatted_content.split('"')
                    alt_parts = []
                    for i, part in enumerate(parts):
                        alt_parts.append(part)
                        if i < len(parts) - 1:
                            alt_parts.append('«' if i % 2 == 0 else '»')
                    formatted_content = "".join(alt_parts)
                    
                    # 2. Spaces around dash
                    formatted_content = re.sub(r'\s+-\s+', ' — ', formatted_content)
                    
                    # 3. Clean spaces, preserving indent
                    lines = formatted_content.splitlines()
                    cleaned_lines = []
                    for line in lines:
                        leading_spaces = len(line) - len(line.lstrip(' '))
                        content_part = line.lstrip(' ')
                        content_part = re.sub(r' {2,}', ' ', content_part)
                        cleaned_lines.append(' ' * leading_spaces + content_part)
                    formatted_content = '\n'.join(cleaned_lines)
                    
                    # 4. Paragraph separation max 1 empty line
                    formatted_content = re.sub(r'\n{3,}', '\n\n', formatted_content)
                    
                    # 5. Clean Markdown headers
                    header_lines = formatted_content.splitlines()
                    cleaned_headers = []
                    for line in header_lines:
                        match = re.match(r'^(#+)\s*(.*)$', line)
                        if match:
                            hashes = match.group(1)
                            h_content = match.group(2).strip()
                            if h_content:
                                h_content = h_content[0].upper() + h_content[1:]
                            cleaned_headers.append(f"{hashes} {h_content}")
                        else:
                            cleaned_headers.append(line)
                    formatted_content = '\n'.join(cleaned_headers)
                
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(formatted_content)
                set_file_permissions_and_ownership(file_path)
                    
                self.send_response(200)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "message": "Файл успешно отформатирован"}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.end_headers()
                self.wfile.write(json.dumps({"error": f"Ошибка при форматировании файла: {str(e)}"}).encode('utf-8'))
        else:
            self.send_response(404)
            self.end_headers()

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)
    
    socketserver.ThreadingTCPServer.allow_reuse_address = True
    
    with socketserver.ThreadingTCPServer(("", PORT), NoCacheProxyHTTPRequestHandler) as httpd:
        print(f"Lumina Chat Server running at: http://localhost:{PORT}")
        print("Press Ctrl+C to stop the server.")
        try:
            httpd.serve_forever()
        except KeyboardInterrupt:
            print("\nStopping server...")
            httpd.server_close()
